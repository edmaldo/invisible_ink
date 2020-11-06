import docx
from docx.shared import RGBColor, Pt

fake_text = docx.Document('docs/fakeMessage.docx')
fake_list = []
for paragraph in fake_text.paragraphs:
    fake_list.append(paragraph.text)

real_text = docx.Document('docs/vigenereMessage.docx')
real_list = []
for paragraph in real_text.paragraphs:
    if len(paragraph.text) != 0:
        real_list.append(paragraph.text)

doc = docx.Document('docs/emptyTemplate.docx')

doc.add_heading('Morland Holmes', 0)
subtitle = doc.add_heading('Global Consulting & Negotiations', 1)
subtitle.alignment = 1
doc.add_heading('', 1)
doc.add_paragraph('July 21, 2020')
doc.add_paragraph('')

def set_spacing(paragraph):
    """Use docx to set line spacing between paragraphs."""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

length_real = len(real_list)
count_real = 0

for line in fake_list:
    if count_real < length_real and line == "":
        paragraph = doc.add_paragraph(real_list[count_real])
        paragraph_index = len(doc.paragraphs) - 1
        run = doc.paragraphs[paragraph_index].runs[0]
        font = run.font
        font.color.rgb = RGBColor(255, 255, 255)
        count_real += 1
    else:
        paragraph = doc.add_paragraph(line)

    set_spacing(paragraph)

doc.save('ciphered_message.docx')

print("Done")
